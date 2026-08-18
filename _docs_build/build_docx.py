# -*- coding: utf-8 -*-
# =============================================================================
#  Coarse-Graining Genesis Framework V4.0
#
#  Author:      Jinku Guo <guojk@nwpu.edu.cn>
#  Affiliation: Northwestern Polytechnical University, Xi'an 710072, China
#
#  Part of the V4 spectral framework, whose physics is presented in the
#  companion papers:
#    [I]  "The spectrum of a compact internal space.
#          I. Gauge structure and fermion content"
#    [II] "The spectrum of a compact internal space.
#          II. Effective couplings and mass scales"
# =============================================================================

"""Markdown -> docx converter (V4 complete guide only).

Supported markup (lightweight, sufficient):
  # / ## / ### / ####     heading (level 1~4)
  ---                     horizontal rule
  plain line              paragraph
  **bold**  *italic*  `code`   inline format
  ~x~ subscript   ^x^ superscript   Word native subscript/superscript
  - item                  unordered list (nested with two spaces)
  1. item                 ordered list
  | a | b |               table (first row is the header)
  ![title](figures/x.png) image
  [[PARAMS:module]]       parameter-table placeholder (generated from params_export.json)
  ``` ... ```             code block (monospace, grey background)
"""
import json
import re
import os

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = os.path.dirname(os.path.abspath(__file__))
FIG = os.path.join(BASE, "figures")
PARAMS = json.load(open(os.path.join(BASE, "params_export.json"), encoding="utf-8"))

FONT_BODY = "Calibri"
FONT_MONO = "Consolas"

COLOR = {
    "h1": (0x1F, 0x4E, 0x79),
    "h2": (0x1F, 0x4E, 0x79),
    "h3": (0x2E, 0x6B, 0x3A),
    "h4": (0x8A, 0x6D, 0x1A),
    "code": (0x8C, 0x2F, 0x2F),
}


def set_font(run, size=None, bold=None, italic=None, color=None, mono=False, sub=False, sup=False):
    run.font.name = FONT_MONO if mono else FONT_BODY
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), FONT_MONO if mono else FONT_BODY)
    rFonts.set(qn("w:hAnsi"), FONT_MONO if mono else FONT_BODY)
    rFonts.set(qn("w:eastAsia"), FONT_BODY)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if italic is not None:
        run.font.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor(*color)
    if sub:
        run.font.subscript = True
    if sup:
        run.font.superscript = True


def add_inline(p, text, base_size=11):
    """Parse inline markup and write to the paragraph: **bold** *italic* `code` ~sub~ ^sup^"""
    token = re.compile(r"(\*\*.+?\*\*|\*.+?\*|`.+?`|~.+?~|\^.+?\^)")
    parts = token.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            r = p.add_run(part[2:-2]); set_font(r, size=base_size, bold=True)
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            r = p.add_run(part[1:-1]); set_font(r, size=base_size, italic=True)
        elif part.startswith("`") and part.endswith("`") and len(part) > 2:
            r = p.add_run(part[1:-1]); set_font(r, size=base_size - 1, mono=True, color=COLOR["code"])
        elif part.startswith("~") and part.endswith("~") and len(part) > 2:
            r = p.add_run(part[1:-1]); set_font(r, size=base_size, sub=True)
        elif part.startswith("^") and part.endswith("^") and len(part) > 2:
            r = p.add_run(part[1:-1]); set_font(r, size=base_size, sup=True)
        else:
            r = p.add_run(part); set_font(r, size=base_size)
    return p


def add_params_table(doc, module_name):
    """Generate the parameter table for a module name (matched from params_export.json by writer)."""
    rows = []
    for key, v in PARAMS.items():
        w = v.get("writer", "")
        if module_name in w:
            rows.append((key, v))
    if not rows:
        return
    # table header
    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(["Parameter", "Value / content", "Derivation and precision (note)"]):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        r = p.add_run(h); set_font(r, size=9.5, bold=True, color=(0xFF, 0xFF, 0xFF))
        shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), "1F4E79")
        hdr[i]._tc.get_or_add_tcPr().append(shd)
    for key, v in rows:
        cells = table.add_row().cells
        # parameter name + role
        p = cells[0].paragraphs[0]
        r = p.add_run(key); set_font(r, size=9, bold=True, mono=True)
        role = v.get("role", "")
        if role:
            p2 = cells[0].add_paragraph()
            r2 = p2.add_run(f"[{role}]"); set_font(r2, size=8, color=(0x88, 0x88, 0x88))
        # value
        val = v.get("value")
        pv = cells[1].paragraphs[0]
        if isinstance(val, dict):
            txt = "; ".join(f"{k}={vv}" for k, vv in val.items())
        else:
            txt = repr(val)
        rv = pv.add_run(txt); set_font(rv, size=8.5, mono=True)
        # note
        pn = cells[2].paragraphs[0]
        rn = pn.add_run(v.get("note", "")); set_font(rn, size=8.5)
    doc.add_paragraph()


def add_toc(doc):
    """Insert the automatic table of contents (TOC field); open in Word and press F9 to generate."""
    p = doc.add_paragraph()
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "(Contents: open, select all with Ctrl+A, then press F9 to update the field and generate the full table of contents)"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r = run._r
    r.append(fld_begin); r.append(instr); r.append(fld_sep)
    r.append(placeholder); r.append(fld_end)


def render_md(doc, md_text):
    lines = md_text.split("\n")
    i = 0
    in_code = False
    code_buf = []
    list_stack = []
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # skip HTML comments (the author/affiliation banner)
        if stripped.startswith("<!--"):
            while i < len(lines) and "-->" not in lines[i]:
                i += 1
            i += 1
            continue

        # code block
        if stripped.startswith("```"):
            if in_code:
                # end code block
                for cl in code_buf:
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.3)
                    p.paragraph_format.space_after = Pt(0)
                    r = p.add_run(cl); set_font(r, size=8.5, mono=True, color=COLOR["code"])
                in_code = False
                code_buf = []
            else:
                in_code = True
                code_buf = []
            i += 1
            continue
        if in_code:
            code_buf.append(line)
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        # heading
        m = re.match(r"^(#{1,4})\s+(.*)$", stripped)
        if m:
            level = len(m.group(1))
            text = m.group(2)
            p = doc.add_heading("", level=level)
            sizes = {1: 20, 2: 16, 3: 13, 4: 11.5}
            r = p.add_run(text)
            set_font(r, size=sizes[level], bold=True, color=COLOR[f"h{level}"])
            if level == 1:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
            p.paragraph_format.space_after = Pt(6)
            i += 1
            continue

        # horizontal rule
        if stripped == "---":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:space"), "1"); bottom.set(qn("w:color"), "AAAAAA")
            pBdr.append(bottom); pPr.append(pBdr)
            i += 1
            continue

        # image
        m = re.match(r"^!\[(.*?)\]\((.*?)\)\s*$", stripped)
        if m:
            title, path = m.group(1), m.group(2)
            full = os.path.join(BASE, path) if not os.path.isabs(path) else path
            if os.path.exists(full):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(full, width=Inches(6.3))
                if title:
                    cap = doc.add_paragraph()
                    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    r = cap.add_run(title)
                    set_font(r, size=9, italic=True, color=(0x66, 0x66, 0x66))
            i += 1
            continue

        # automatic table of contents
        if stripped == "[[TOC]]":
            add_toc(doc)
            i += 1
            continue

        # parameter-table placeholder
        m = re.match(r"^\[\[PARAMS:(.*?)\]\]\s*$", stripped)
        if m:
            add_params_table(doc, m.group(1).strip())
            i += 1
            continue

        # table
        if stripped.startswith("|") and stripped.endswith("|"):
            # collect consecutive table rows
            tbl_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                tbl_lines.append(lines[i].strip())
                i += 1
            # filter separator rows
            data_rows = []
            for tl in tbl_lines:
                cells = [c.strip() for c in tl.strip("|").split("|")]
                if all(re.fullmatch(r":?-{2,}:?", c) for c in cells):
                    continue
                data_rows.append(cells)
            if data_rows:
                ncols = max(len(r) for r in data_rows)
                table = doc.add_table(rows=len(data_rows), cols=ncols)
                table.style = "Light Grid Accent 1"
                for ri, row in enumerate(data_rows):
                    for ci in range(ncols):
                        cell_text = row[ci] if ci < len(row) else ""
                        cell = table.rows[ri].cells[ci]
                        cell.text = ""
                        p = cell.paragraphs[0]
                        if ri == 0:
                            add_inline(p, cell_text, base_size=9.5)
                            for r in p.runs:
                                r.font.bold = True
                        else:
                            add_inline(p, cell_text, base_size=9.5)
                doc.add_paragraph()
            continue

        # list
        m = re.match(r"^(\s*)([-*]|\d+\.)\s+(.*)$", line)
        if m:
            indent = len(m.group(1))
            ordered = m.group(2)[0].isdigit()
            text = m.group(3)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25 + indent * 0.1)
            p.paragraph_format.space_after = Pt(2)
            prefix = "• " if not ordered else f"{m.group(2)} "
            r = p.add_run(prefix)
            set_font(r, size=11, bold=True, color=(0x1F, 0x4E, 0x79))
            add_inline(p, text, base_size=11)
            i += 1
            continue

        # plain paragraph
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        add_inline(p, stripped, base_size=11)
        i += 1
    return doc


def main():
    doc = Document()
    # page and default font
    st = doc.styles["Normal"]
    st.font.name = FONT_BODY
    st.font.size = Pt(11)
    st.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)

    # read content files (in order)
    parts = ["part0_cover.md", "part0_intro.md", "part1_symmetry.md", "part2_params.md", "part3_supplement.md"]
    for fname in parts:
        path = os.path.join(BASE, fname)
        if os.path.exists(path):
            md = open(path, encoding="utf-8").read()
            render_md(doc, md)

    out = os.path.join(BASE, "..", "docs", "V4_COMPLETE_GUIDE.docx")
    doc.save(out)
    print("Generated:", os.path.abspath(out))


if __name__ == "__main__":
    main()
